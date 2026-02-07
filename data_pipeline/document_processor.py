"""
Document Processor for EACP Data Pipeline
Handles multi-format document processing for RAG and fine-tuning
"""

from typing import Dict, Any, Optional, List
from datetime import datetime
import logging
import os

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Processes documents from various formats into structured text.
    Supports PDF, DOCX, HTML, Markdown, CSV, and plain text.
    """

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.supported_formats = [".txt", ".md", ".pdf", ".docx", ".html", ".csv", ".json", ".jsonl"]
        self.processed_count = 0

    def process_file(self, filepath: str) -> Dict[str, Any]:
        """Process a single file and extract text content"""
        if not os.path.exists(filepath):
            return {"error": f"File not found: {filepath}", "success": False}

        ext = os.path.splitext(filepath)[1].lower()

        if ext not in self.supported_formats:
            return {"error": f"Unsupported format: {ext}", "success": False}

        try:
            if ext in (".txt", ".md"):
                result = self._process_text(filepath)
            elif ext == ".pdf":
                result = self._process_pdf(filepath)
            elif ext == ".docx":
                result = self._process_docx(filepath)
            elif ext == ".html":
                result = self._process_html(filepath)
            elif ext == ".csv":
                result = self._process_csv(filepath)
            elif ext in (".json", ".jsonl"):
                result = self._process_json(filepath)
            else:
                result = self._process_text(filepath)

            self.processed_count += 1
            result["success"] = True
            result["processed_at"] = datetime.now().isoformat()
            return result

        except Exception as e:
            logger.error(f"Error processing {filepath}: {e}")
            return {"error": str(e), "success": False, "filepath": filepath}

    def process_directory(self, dirpath: str,
                          extensions: Optional[List[str]] = None,
                          recursive: bool = True) -> List[Dict[str, Any]]:
        """Process all documents in a directory"""
        extensions = extensions or self.supported_formats
        results = []

        walker = os.walk(dirpath) if recursive else [(dirpath, [], os.listdir(dirpath))]

        for root, _, files in walker:
            for filename in files:
                ext = os.path.splitext(filename)[1].lower()
                if ext in extensions:
                    filepath = os.path.join(root, filename)
                    result = self.process_file(filepath)
                    results.append(result)

        logger.info(f"Processed {len(results)} documents from {dirpath}")
        return results

    def _process_text(self, filepath: str) -> Dict[str, Any]:
        """Process plain text or markdown file"""
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()

        return {
            "filepath": filepath,
            "format": os.path.splitext(filepath)[1][1:],
            "content": content,
            "word_count": len(content.split()),
            "char_count": len(content),
            "line_count": content.count("\n") + 1
        }

    def _process_pdf(self, filepath: str) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            import pdfplumber

            pages = []
            full_text = []

            with pdfplumber.open(filepath) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    tables = page.extract_tables() or []

                    page_data = {
                        "page_number": i + 1,
                        "text": text,
                        "tables": [
                            [[cell or "" for cell in row] for row in table]
                            for table in tables
                        ],
                        "has_tables": len(tables) > 0
                    }
                    pages.append(page_data)
                    full_text.append(text)

            content = "\n\n".join(full_text)

            return {
                "filepath": filepath,
                "format": "pdf",
                "content": content,
                "pages": pages,
                "total_pages": len(pages),
                "word_count": len(content.split()),
                "has_tables": any(p["has_tables"] for p in pages)
            }

        except ImportError:
            logger.warning("pdfplumber not available, falling back to basic read")
            return {
                "filepath": filepath,
                "format": "pdf",
                "content": "[PDF processing requires pdfplumber]",
                "error": "pdfplumber not installed"
            }

    def _process_docx(self, filepath: str) -> Dict[str, Any]:
        """Process Word document"""
        try:
            import docx

            doc = docx.Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)

            return {
                "filepath": filepath,
                "format": "docx",
                "content": content,
                "paragraphs": len(paragraphs),
                "tables": tables,
                "word_count": len(content.split())
            }

        except ImportError:
            return {
                "filepath": filepath,
                "format": "docx",
                "content": "[DOCX processing requires python-docx]",
                "error": "python-docx not installed"
            }

    def _process_html(self, filepath: str) -> Dict[str, Any]:
        """Process HTML file"""
        try:
            from bs4 import BeautifulSoup

            with open(filepath, "r", encoding="utf-8") as f:
                html = f.read()

            soup = BeautifulSoup(html, "html.parser")

            # Remove script and style
            for elem in soup(["script", "style", "nav", "footer", "header"]):
                elem.decompose()

            text = soup.get_text(separator="\n", strip=True)
            title = soup.title.string if soup.title else ""

            # Extract links
            links = [{"text": a.text.strip(), "href": a.get("href", "")}
                     for a in soup.find_all("a") if a.get("href")]

            return {
                "filepath": filepath,
                "format": "html",
                "content": text,
                "title": title,
                "links": links[:50],  # Limit links
                "word_count": len(text.split())
            }

        except ImportError:
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
            return {
                "filepath": filepath,
                "format": "html",
                "content": content,
                "word_count": len(content.split())
            }

    def _process_csv(self, filepath: str) -> Dict[str, Any]:
        """Process CSV file"""
        try:
            import pandas as pd

            df = pd.read_csv(filepath)

            # Convert rows to text
            rows_text = []
            for _, row in df.iterrows():
                row_text = " | ".join(f"{col}: {val}" for col, val in row.items())
                rows_text.append(row_text)

            content = "\n".join(rows_text)

            return {
                "filepath": filepath,
                "format": "csv",
                "content": content,
                "columns": list(df.columns),
                "row_count": len(df),
                "word_count": len(content.split()),
                "sample_rows": df.head(5).to_dict(orient="records")
            }

        except ImportError:
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
            return {
                "filepath": filepath,
                "format": "csv",
                "content": content,
                "word_count": len(content.split())
            }

    def _process_json(self, filepath: str) -> Dict[str, Any]:
        """Process JSON/JSONL file"""
        import json

        ext = os.path.splitext(filepath)[1].lower()

        if ext == ".jsonl":
            records = []
            with open(filepath, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line:
                        records.append(json.loads(line))

            content = "\n".join(json.dumps(r, ensure_ascii=False) for r in records)

            return {
                "filepath": filepath,
                "format": "jsonl",
                "content": content,
                "record_count": len(records),
                "sample_records": records[:5],
                "word_count": len(content.split())
            }
        else:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)

            content = json.dumps(data, indent=2, ensure_ascii=False)

            return {
                "filepath": filepath,
                "format": "json",
                "content": content,
                "is_array": isinstance(data, list),
                "record_count": len(data) if isinstance(data, list) else 1,
                "word_count": len(content.split())
            }

    def get_stats(self) -> Dict[str, Any]:
        """Get processing statistics"""
        return {
            "total_processed": self.processed_count,
            "supported_formats": self.supported_formats
        }
