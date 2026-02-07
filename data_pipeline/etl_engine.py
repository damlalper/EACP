"""
ETL Engine for EACP
Manages data extraction, transformation, and loading pipelines
for RAG ingestion and fine-tuning dataset preparation
"""

from typing import Dict, Any, Optional, List, Callable
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
import logging
import json
import os

logger = logging.getLogger(__name__)


class JobStatus(str, Enum):
    PENDING = "pending"
    RUNNING = "running"
    COMPLETED = "completed"
    FAILED = "failed"
    CANCELLED = "cancelled"


class DataSource(str, Enum):
    FILE = "file"           # Local files (PDF, DOCX, TXT, CSV, JSON)
    DIRECTORY = "directory"  # Directory of files
    URL = "url"             # Web URLs
    DATABASE = "database"   # Database connections
    API = "api"             # External APIs
    S3 = "s3"              # AWS S3 buckets
    JSONL = "jsonl"         # JSONL files (fine-tuning format)


@dataclass
class ETLJob:
    """Represents an ETL pipeline job"""
    job_id: str
    name: str
    source_type: str
    source_config: Dict[str, Any]
    target: str  # "rag" or "fine_tuning"
    status: str = JobStatus.PENDING
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    records_processed: int = 0
    records_failed: int = 0
    errors: List[str] = field(default_factory=list)
    output_path: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class ETLEngine:
    """
    ETL Engine for processing data into RAG-ready or fine-tuning-ready formats.

    Supported input formats:
    - Text files (.txt, .md)
    - PDF documents (.pdf)
    - JSON/JSONL files (.json, .jsonl)
    - CSV files (.csv)
    - Word documents (.docx)
    - Web pages (URL extraction)

    Output formats:
    - RAG: Chunked documents with metadata for vector DB ingestion
    - Fine-tuning: Instruction/input/output JSONL format
    """

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.jobs: Dict[str, ETLJob] = {}
        self.transformers: Dict[str, Callable] = {}
        self._register_default_transformers()

    def _register_default_transformers(self):
        """Register built-in data transformers"""
        self.transformers["text_chunker"] = self._chunk_text
        self.transformers["json_flattener"] = self._flatten_json
        self.transformers["csv_to_documents"] = self._csv_to_documents
        self.transformers["instruction_formatter"] = self._format_for_fine_tuning
        self.transformers["metadata_enricher"] = self._enrich_metadata
        self.transformers["text_cleaner"] = self._clean_text

    def register_transformer(self, name: str, func: Callable):
        """Register a custom transformer function"""
        self.transformers[name] = func
        logger.info(f"Registered custom transformer: {name}")

    def create_job(self, name: str, source_type: str,
                   source_config: Dict[str, Any],
                   target: str = "rag") -> ETLJob:
        """Create a new ETL job"""
        job_id = f"etl_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{len(self.jobs)}"

        job = ETLJob(
            job_id=job_id,
            name=name,
            source_type=source_type,
            source_config=source_config,
            target=target
        )

        self.jobs[job_id] = job
        logger.info(f"Created ETL job: {job_id} ({name})")
        return job

    def run_job(self, job_id: str) -> ETLJob:
        """Execute an ETL job"""
        job = self.jobs.get(job_id)
        if not job:
            raise ValueError(f"Job not found: {job_id}")

        job.status = JobStatus.RUNNING
        job.started_at = datetime.now().isoformat()
        logger.info(f"Running ETL job: {job_id}")

        try:
            # Extract
            raw_data = self._extract(job)
            logger.info(f"Extracted {len(raw_data)} records")

            # Transform
            transformed = self._transform(job, raw_data)
            logger.info(f"Transformed {len(transformed)} records")

            # Load
            output_path = self._load(job, transformed)

            job.status = JobStatus.COMPLETED
            job.completed_at = datetime.now().isoformat()
            job.records_processed = len(transformed)
            job.output_path = output_path
            logger.info(f"ETL job completed: {job_id}, {len(transformed)} records")

        except Exception as e:
            job.status = JobStatus.FAILED
            job.completed_at = datetime.now().isoformat()
            job.errors.append(str(e))
            logger.error(f"ETL job failed: {job_id}: {e}")

        return job

    def _extract(self, job: ETLJob) -> List[Dict[str, Any]]:
        """Extract data from source"""
        source_type = job.source_type
        config = job.source_config

        if source_type == DataSource.FILE:
            return self._extract_file(config.get("path", ""))

        elif source_type == DataSource.DIRECTORY:
            return self._extract_directory(
                config.get("path", ""),
                config.get("extensions", [".txt", ".md", ".pdf", ".json"])
            )

        elif source_type == DataSource.JSONL:
            return self._extract_jsonl(config.get("path", ""))

        elif source_type == DataSource.URL:
            return self._extract_url(config.get("urls", []))

        elif source_type == DataSource.API:
            return self._extract_api(config)

        else:
            raise ValueError(f"Unsupported source type: {source_type}")

    def _extract_file(self, filepath: str) -> List[Dict[str, Any]]:
        """Extract data from a single file"""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        ext = os.path.splitext(filepath)[1].lower()
        records = []

        if ext in (".txt", ".md"):
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
            records.append({
                "content": content,
                "source": filepath,
                "type": ext[1:],
                "size": len(content)
            })

        elif ext == ".json":
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                for i, item in enumerate(data):
                    records.append({
                        "content": json.dumps(item, ensure_ascii=False),
                        "source": filepath,
                        "type": "json",
                        "index": i
                    })
            else:
                records.append({
                    "content": json.dumps(data, ensure_ascii=False),
                    "source": filepath,
                    "type": "json"
                })

        elif ext == ".csv":
            try:
                import pandas as pd
                df = pd.read_csv(filepath)
                for i, row in df.iterrows():
                    records.append({
                        "content": row.to_json(),
                        "source": filepath,
                        "type": "csv",
                        "index": i,
                        "columns": list(df.columns)
                    })
            except ImportError:
                logger.warning("pandas not available for CSV parsing")
                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()
                records.append({"content": content, "source": filepath, "type": "csv"})

        elif ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text() or ""
                        if text.strip():
                            records.append({
                                "content": text,
                                "source": filepath,
                                "type": "pdf",
                                "page": i + 1,
                                "total_pages": len(pdf.pages)
                            })
            except ImportError:
                logger.warning("pdfplumber not available for PDF parsing")

        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(filepath)
                full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                records.append({
                    "content": full_text,
                    "source": filepath,
                    "type": "docx",
                    "paragraphs": len(doc.paragraphs)
                })
            except ImportError:
                logger.warning("python-docx not available for DOCX parsing")

        return records

    def _extract_directory(self, dirpath: str,
                           extensions: List[str]) -> List[Dict[str, Any]]:
        """Extract data from all matching files in a directory"""
        records = []

        if not os.path.isdir(dirpath):
            raise NotADirectoryError(f"Directory not found: {dirpath}")

        for root, _, files in os.walk(dirpath):
            for filename in files:
                ext = os.path.splitext(filename)[1].lower()
                if ext in extensions:
                    filepath = os.path.join(root, filename)
                    try:
                        file_records = self._extract_file(filepath)
                        records.extend(file_records)
                    except Exception as e:
                        logger.warning(f"Failed to extract {filepath}: {e}")

        return records

    def _extract_jsonl(self, filepath: str) -> List[Dict[str, Any]]:
        """Extract data from JSONL file (common fine-tuning format)"""
        records = []

        with open(filepath, "r", encoding="utf-8") as f:
            for i, line in enumerate(f):
                line = line.strip()
                if line:
                    try:
                        data = json.loads(line)
                        data["_source"] = filepath
                        data["_index"] = i
                        records.append(data)
                    except json.JSONDecodeError:
                        logger.warning(f"Invalid JSON at line {i + 1}")

        return records

    def _extract_url(self, urls: List[str]) -> List[Dict[str, Any]]:
        """Extract content from URLs"""
        records = []

        try:
            import requests
            from bs4 import BeautifulSoup

            for url in urls:
                try:
                    response = requests.get(url, timeout=30)
                    response.raise_for_status()

                    soup = BeautifulSoup(response.text, "html.parser")
                    # Remove script and style elements
                    for elem in soup(["script", "style", "nav", "footer"]):
                        elem.decompose()

                    text = soup.get_text(separator="\n", strip=True)
                    records.append({
                        "content": text,
                        "source": url,
                        "type": "url",
                        "title": soup.title.string if soup.title else ""
                    })
                except Exception as e:
                    logger.warning(f"Failed to extract URL {url}: {e}")

        except ImportError:
            logger.warning("requests/beautifulsoup4 not available for URL extraction")

        return records

    def _extract_api(self, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract data from an API endpoint"""
        try:
            import requests

            url = config.get("endpoint", "")
            headers = config.get("headers", {})
            params = config.get("params", {})
            method = config.get("method", "GET").upper()

            if method == "GET":
                response = requests.get(url, headers=headers, params=params, timeout=30)
            else:
                response = requests.post(url, headers=headers, json=params, timeout=30)

            response.raise_for_status()
            data = response.json()

            # Extract records from response
            data_key = config.get("data_key")  # e.g., "results", "data", "items"
            if data_key and isinstance(data, dict):
                items = data.get(data_key, [])
            elif isinstance(data, list):
                items = data
            else:
                items = [data]

            return [
                {"content": json.dumps(item, ensure_ascii=False), "source": url, "type": "api"}
                for item in items
            ]

        except ImportError:
            logger.warning("requests not available for API extraction")
            return []

    def _transform(self, job: ETLJob, data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Apply transformations based on target format"""
        if job.target == "rag":
            return self._transform_for_rag(data, job.source_config)
        elif job.target == "fine_tuning":
            return self._transform_for_fine_tuning(data, job.source_config)
        else:
            return data

    def _transform_for_rag(self, data: List[Dict[str, Any]],
                            config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Transform data for RAG vector DB ingestion"""
        chunk_size = config.get("chunk_size", 500)
        chunk_overlap = config.get("chunk_overlap", 50)
        documents = []

        for record in data:
            content = record.get("content", "")
            if not content.strip():
                continue

            # Clean text
            content = self._clean_text(content)

            # Chunk text
            chunks = self._chunk_text(content, chunk_size=chunk_size, overlap=chunk_overlap)

            for i, chunk in enumerate(chunks):
                doc = {
                    "text": chunk,
                    "metadata": {
                        "source": record.get("source", "unknown"),
                        "type": record.get("type", "unknown"),
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "page": record.get("page"),
                        "title": record.get("title", ""),
                        "processed_at": datetime.now().isoformat()
                    }
                }
                documents.append(doc)

        return documents

    def _transform_for_fine_tuning(self, data: List[Dict[str, Any]],
                                    config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Transform data for fine-tuning (instruction/input/output format)"""
        format_type = config.get("format", "alpaca")  # alpaca, sharegpt, openai
        training_data = []

        for record in data:
            # If already in instruction format
            if "instruction" in record:
                entry = self._format_for_fine_tuning(record, format_type)
                if entry:
                    training_data.append(entry)
                continue

            # Convert content to Q&A pairs using simple heuristics
            content = record.get("content", "")
            if not content.strip():
                continue

            # Create a summarization training example
            entry = {
                "instruction": "Summarize the following text.",
                "input": content[:1000],  # Limit input length
                "output": content[:200] + "..."  # Placeholder summary
            }

            formatted = self._format_for_fine_tuning(entry, format_type)
            if formatted:
                training_data.append(formatted)

        return training_data

    def _load(self, job: ETLJob, data: List[Dict[str, Any]]) -> str:
        """Load transformed data to target"""
        output_dir = self.config.get("output_dir", "data/processed")
        os.makedirs(output_dir, exist_ok=True)

        if job.target == "rag":
            output_path = os.path.join(output_dir, f"{job.job_id}_rag.json")
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

        elif job.target == "fine_tuning":
            output_path = os.path.join(output_dir, f"{job.job_id}_training.jsonl")
            with open(output_path, "w", encoding="utf-8") as f:
                for record in data:
                    f.write(json.dumps(record, ensure_ascii=False) + "\n")

        else:
            output_path = os.path.join(output_dir, f"{job.job_id}_output.json")
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

        return output_path

    # ==================== Transformer Functions ====================

    def _chunk_text(self, text: str, chunk_size: int = 500,
                    overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks"""
        if isinstance(text, dict):
            text = text.get("content", str(text))

        words = text.split()
        chunks = []

        if len(words) <= chunk_size:
            return [text]

        i = 0
        while i < len(words):
            chunk = " ".join(words[i:i + chunk_size])
            chunks.append(chunk)
            i += chunk_size - overlap

        return chunks

    def _flatten_json(self, data: Any, prefix: str = "") -> Dict[str, Any]:
        """Flatten nested JSON into flat key-value pairs"""
        flat = {}

        if isinstance(data, dict):
            for key, value in data.items():
                new_key = f"{prefix}.{key}" if prefix else key
                if isinstance(value, (dict, list)):
                    flat.update(self._flatten_json(value, new_key))
                else:
                    flat[new_key] = value
        elif isinstance(data, list):
            for i, item in enumerate(data):
                new_key = f"{prefix}[{i}]"
                if isinstance(item, (dict, list)):
                    flat.update(self._flatten_json(item, new_key))
                else:
                    flat[new_key] = item

        return flat

    def _csv_to_documents(self, records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Convert CSV records to document format"""
        documents = []
        for record in records:
            content = record.get("content", "")
            try:
                data = json.loads(content) if isinstance(content, str) else content
                text = " | ".join(f"{k}: {v}" for k, v in data.items())
                documents.append({"text": text, "metadata": record})
            except (json.JSONDecodeError, AttributeError):
                documents.append({"text": str(content), "metadata": record})
        return documents

    def _format_for_fine_tuning(self, record: Dict[str, Any],
                                 format_type: str = "alpaca") -> Optional[Dict[str, Any]]:
        """Format a record for fine-tuning"""
        if format_type == "alpaca":
            return {
                "instruction": record.get("instruction", ""),
                "input": record.get("input", ""),
                "output": record.get("output", "")
            }
        elif format_type == "sharegpt":
            return {
                "conversations": [
                    {"from": "human", "value": record.get("instruction", "") +
                     ("\n" + record.get("input", "") if record.get("input") else "")},
                    {"from": "gpt", "value": record.get("output", "")}
                ]
            }
        elif format_type == "openai":
            messages = [{"role": "system", "content": "You are a helpful assistant."}]
            user_msg = record.get("instruction", "")
            if record.get("input"):
                user_msg += f"\n\n{record['input']}"
            messages.append({"role": "user", "content": user_msg})
            messages.append({"role": "assistant", "content": record.get("output", "")})
            return {"messages": messages}
        else:
            return record

    def _enrich_metadata(self, record: Dict[str, Any]) -> Dict[str, Any]:
        """Enrich a record with additional metadata"""
        content = record.get("content", record.get("text", ""))
        record["metadata"] = record.get("metadata", {})
        record["metadata"]["word_count"] = len(content.split())
        record["metadata"]["char_count"] = len(content)
        record["metadata"]["enriched_at"] = datetime.now().isoformat()
        return record

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        if isinstance(text, dict):
            text = text.get("content", str(text))

        # Remove excessive whitespace
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line:
                # Collapse multiple spaces
                while "  " in line:
                    line = line.replace("  ", " ")
                cleaned_lines.append(line)

        return "\n".join(cleaned_lines)

    # ==================== Job Management ====================

    def get_job(self, job_id: str) -> Optional[ETLJob]:
        """Get job by ID"""
        return self.jobs.get(job_id)

    def get_all_jobs(self) -> List[ETLJob]:
        """Get all jobs"""
        return list(self.jobs.values())

    def get_job_status(self, job_id: str) -> Dict[str, Any]:
        """Get job status"""
        job = self.jobs.get(job_id)
        if not job:
            return {"error": f"Job not found: {job_id}"}

        return {
            "job_id": job.job_id,
            "name": job.name,
            "status": job.status,
            "records_processed": job.records_processed,
            "records_failed": job.records_failed,
            "output_path": job.output_path,
            "errors": job.errors,
            "created_at": job.created_at,
            "started_at": job.started_at,
            "completed_at": job.completed_at
        }
